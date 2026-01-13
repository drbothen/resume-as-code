"""DOCX provider using python-docx."""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from resume_as_code.models.certification import Certification
from resume_as_code.models.education import Education
from resume_as_code.models.errors import RenderError
from resume_as_code.models.resume import ResumeData, ResumeItem


class DOCXProvider:
    """Provider for generating DOCX resumes using python-docx."""

    def render(self, resume: ResumeData, output_path: Path) -> Path:
        """Render resume to DOCX file.

        Args:
            resume: ResumeData to render.
            output_path: Path for output DOCX file.

        Returns:
            Path to generated DOCX.

        Raises:
            RenderError: If DOCX generation fails.
        """
        try:
            doc = self._build_document(resume)

            # Ensure output directory exists (idempotent for build command's mkdir)
            output_path.parent.mkdir(parents=True, exist_ok=True)

            # Save document
            doc.save(str(output_path))

            return output_path
        except PermissionError as e:
            raise RenderError(
                message=f"Permission denied writing to {output_path}: {e}",
                path=str(output_path),
                suggestion="Check file permissions or close the file if open elsewhere",
            ) from e
        except OSError as e:
            raise RenderError(
                message=f"DOCX generation failed: {e}",
                path=str(output_path),
                suggestion="Ensure the output directory is writable",
            ) from e
        except Exception as e:
            raise RenderError(
                message=f"Failed to render DOCX: {e}",
                path=str(output_path),
                suggestion="Check the resume data for issues",
            ) from e

    def render_to_bytes(self, resume: ResumeData) -> bytes:
        """Render resume to DOCX bytes.

        Useful for streaming or in-memory processing.

        Args:
            resume: ResumeData to render.

        Returns:
            DOCX content as bytes.

        Raises:
            RenderError: If DOCX generation fails.
        """
        try:
            doc = self._build_document(resume)

            # Write to in-memory buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return buffer.read()
        except Exception as e:
            raise RenderError(
                message=f"Failed to render DOCX to bytes: {e}",
                suggestion="Check the resume data for issues",
            ) from e

    def _build_document(self, resume: ResumeData) -> Any:
        """Build the Word document from resume data.

        Args:
            resume: ResumeData to render.

        Returns:
            python-docx Document object.
        """
        doc: Any = Document()

        # Set up page margins
        for doc_section in doc.sections:
            doc_section.top_margin = Inches(0.75)
            doc_section.bottom_margin = Inches(0.75)
            doc_section.left_margin = Inches(0.75)
            doc_section.right_margin = Inches(0.75)

        # Header with name and contact info
        self._add_header(doc, resume)

        # Summary section
        if resume.summary:
            self._add_section_heading(doc, "Summary")
            p = doc.add_paragraph(resume.summary)
            p.paragraph_format.space_after = Pt(12)

        # Experience sections
        for resume_section in resume.sections:
            self._add_section_heading(doc, resume_section.title)
            for idx, item in enumerate(resume_section.items):
                is_last = idx == len(resume_section.items) - 1
                self._add_experience_item(doc, item, is_last=is_last)

        # Education section
        displayable_education = [edu for edu in resume.education if edu.display]
        if displayable_education:
            self._add_section_heading(doc, "Education")
            for idx, edu in enumerate(displayable_education):
                is_last = idx == len(displayable_education) - 1
                self._add_education_item(doc, edu, is_last=is_last)

        # Certifications section
        active_certs = resume.get_active_certifications()
        if active_certs:
            self._add_certifications_section(doc, active_certs)

        # Skills section
        if resume.skills:
            self._add_section_heading(doc, "Skills")
            p = doc.add_paragraph(", ".join(resume.skills))
            p.paragraph_format.space_after = Pt(12)

        return doc

    def _add_header(self, doc: Any, resume: ResumeData) -> None:
        """Add header with contact information.

        Args:
            doc: Word document to add header to.
            resume: Resume data containing contact info.
        """
        # Name (centered, bold, large font)
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(resume.contact.name)
        name_run.bold = True
        name_run.font.size = Pt(24)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_after = Pt(6)

        # Contact line (email | phone | location)
        contact_parts: list[str] = []
        if resume.contact.email:
            contact_parts.append(resume.contact.email)
        if resume.contact.phone:
            contact_parts.append(resume.contact.phone)
        if resume.contact.location:
            contact_parts.append(resume.contact.location)

        if contact_parts:
            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.paragraph_format.space_after = Pt(3)

        # Links line (linkedin | github | website)
        link_parts: list[str] = []
        if resume.contact.linkedin:
            link_parts.append(resume.contact.linkedin)
        if resume.contact.github:
            link_parts.append(resume.contact.github)
        if resume.contact.website:
            link_parts.append(resume.contact.website)

        if link_parts:
            links_para = doc.add_paragraph(" | ".join(link_parts))
            links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            links_para.paragraph_format.space_after = Pt(12)

    def _add_section_heading(self, doc: Any, title: str) -> None:
        """Add a section heading (Heading 2 style).

        Args:
            doc: Word document to add heading to.
            title: Section title text.
        """
        heading = doc.add_heading(title, level=2)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

    def _add_experience_item(self, doc: Any, item: ResumeItem, *, is_last: bool = False) -> None:
        """Add an experience item with title, org, dates, and bullets.

        Args:
            doc: Word document to add item to.
            item: Experience item to render.
            is_last: Whether this is the last item in the section.
        """
        # Title line with org and dates
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(item.title)
        title_run.bold = True

        if item.organization:
            title_para.add_run(f" | {item.organization}")

        if item.start_date:
            dates = f"{item.start_date} - {item.end_date or 'Present'}"
            title_para.add_run(f"  ({dates})")

        title_para.paragraph_format.space_after = Pt(3)

        # Scope indicators for executive roles
        if item.scope_team_size or item.scope_budget:
            scope_parts: list[str] = []
            if item.scope_team_size:
                scope_parts.append(f"Team: {item.scope_team_size}")
            if item.scope_budget:
                scope_parts.append(f"Budget: {item.scope_budget}")

            scope_para = doc.add_paragraph(", ".join(scope_parts))
            scope_para.paragraph_format.left_indent = Inches(0.25)
            for run in scope_para.runs:
                run.italic = True
            scope_para.paragraph_format.space_after = Pt(3)

        # Bullets using Word list style
        for bullet in item.bullets:
            bullet_para = doc.add_paragraph(bullet.text, style="List Bullet")
            bullet_para.paragraph_format.space_after = Pt(3)

        # Add spacer only if not the last item (avoid trailing whitespace)
        if not is_last:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)

    def _add_education_item(self, doc: Any, edu: Education, *, is_last: bool = False) -> None:
        """Add an education item.

        Args:
            doc: Word document to add item to.
            edu: Education item to render.
            is_last: Whether this is the last item in the section.
        """
        # Format: "Degree, Institution, Year - Honors" or "(GPA: X)"
        parts = [edu.degree, edu.institution]
        if edu.year:
            parts.append(edu.year)

        text = ", ".join(parts)
        if edu.honors:
            text += f" - {edu.honors}"
        elif edu.gpa:
            text += f" (GPA: {edu.gpa})"

        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(3)

        # Add spacer only if not the last item
        if not is_last:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)

    def _add_certifications_section(self, doc: Any, certifications: list[Certification]) -> None:
        """Add certifications section with Word bullet list formatting.

        Args:
            doc: Word document to add section to.
            certifications: List of active certifications to render.
        """
        self._add_section_heading(doc, "Certifications")

        for cert in certifications:
            # Build certification display text: "Name, Issuer, Year"
            parts: list[str] = [cert.name]
            if cert.issuer:
                parts.append(cert.issuer)
            if cert.date:
                parts.append(cert.date[:4])  # Year only
            if cert.expires:
                parts.append(f"expires {cert.expires[:4]}")

            cert_text = ", ".join(parts)
            bullet_para = doc.add_paragraph(cert_text, style="List Bullet")
            bullet_para.paragraph_format.space_after = Pt(3)
